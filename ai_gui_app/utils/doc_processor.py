import os
try:
    import PyPDF2
except ImportError:
    print("PyPDF2 not installed. PDF processing will not be available.")
    PyPDF2 = None # Placeholder

try:
    import openpyxl
except ImportError:
    print("openpyxl not installed. Excel processing will not be available.")
    openpyxl = None # Placeholder

try:
    from docx import Document as DocxDocument # Alias to avoid conflict if we define our own Document class
except ImportError:
    print("python-docx not installed. Word processing will not be available.")
    DocxDocument = None # Placeholder

class DocumentProcessingError(Exception):
    """Custom exception for document processing errors."""
    pass

def read_pdf(file_path: str) -> str:
    """
    Extracts text from a PDF file.
    Args:
        file_path (str): The path to the PDF file.
    Returns:
        str: The extracted text, or an empty string if an error occurs or PyPDF2 is not available.
    Raises:
        DocumentProcessingError: If the file is not found or is not a PDF.
    """
    if not PyPDF2:
        print("PyPDF2 is not installed. Cannot process PDF.")
        return ""
    if not os.path.exists(file_path):
        raise DocumentProcessingError(f"PDF file not found: {file_path}")
    if not file_path.lower().endswith(".pdf"):
        raise DocumentProcessingError(f"File is not a PDF: {file_path}")

    text = []
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return "\n".join(text)
    except PyPDF2.errors.PdfReadError as e: # More specific exception for PyPDF2
        raise DocumentProcessingError(f"Error reading PDF file {file_path}: {e}")
    except Exception as e:
        # Catch other potential errors during processing
        raise DocumentProcessingError(f"An unexpected error occurred while processing PDF {file_path}: {e}")


def read_excel(file_path: str) -> str:
    """
    Extracts text from an Excel file (all sheets, all cells).
    Args:
        file_path (str): The path to the Excel file (.xlsx).
    Returns:
        str: The concatenated text from all cells in all sheets, or an empty string if an error occurs or openpyxl is not available.
    Raises:
        DocumentProcessingError: If the file is not found or is not an Excel file.
    """
    if not openpyxl:
        print("openpyxl is not installed. Cannot process Excel.")
        return ""
    if not os.path.exists(file_path):
        raise DocumentProcessingError(f"Excel file not found: {file_path}")
    if not file_path.lower().endswith((".xlsx")): # openpyxl primarily supports .xlsx
        raise DocumentProcessingError(f"File is not an XLSX Excel file: {file_path}")

    text = []
    try:
        workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True) # data_only to get cell values not formulas
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        text.append(str(cell.value))
        return "\n".join(text)
    except Exception as e: # openpyxl can raise various exceptions
        raise DocumentProcessingError(f"Error reading Excel file {file_path}: {e}")


def read_word(file_path: str) -> str:
    """
    Extracts text from a Word document (.docx).
    Args:
        file_path (str): The path to the Word file (.docx).
    Returns:
        str: The extracted text, or an empty string if an error occurs or python-docx is not available.
    Raises:
        DocumentProcessingError: If the file is not found or is not a Word file.
    """
    if not DocxDocument:
        print("python-docx is not installed. Cannot process Word document.")
        return ""
    if not os.path.exists(file_path):
        raise DocumentProcessingError(f"Word file not found: {file_path}")
    if not file_path.lower().endswith(".docx"):
        raise DocumentProcessingError(f"File is not a DOCX Word file: {file_path}")

    text = []
    try:
        document = DocxDocument(file_path)
        for para in document.paragraphs:
            text.append(para.text)
        return "\n".join(text)
    except Exception as e: # python-docx can raise various exceptions (e.g., PackageNotFoundError)
        raise DocumentProcessingError(f"Error reading Word file {file_path}: {e}")

# Example Usage (for testing purposes)
if __name__ == '__main__':
    # Create dummy files for testing
    # Note: Actual file creation for .pdf, .xlsx, .docx is complex to do in a simple script.
    # These tests would typically be run in an environment where such files exist.
    print("Document Processor Module")

    # Test PDF (requires a dummy file named test.pdf)
    # For now, we'll just show how it would be called.
    # with open("test.pdf", "wb") as f: # A real PDF cannot be created this simply
    #     f.write(b"%PDF-1.4...")
    try:
        print("\n--- Testing PDF (will fail if test.pdf not present or invalid) ---")
        # pdf_text = read_pdf("test.pdf")
        # print(f"Extracted PDF Text: '{pdf_text[:100]}...'")
        print("PDF test skipped (requires a valid test.pdf file).")
    except DocumentProcessingError as e:
        print(f"PDF Error: {e}")
    except FileNotFoundError:
        print("test.pdf not found for testing.")

    # Test Excel (requires a dummy file named test.xlsx)
    # if openpyxl:
    #     try:
    #         print("\n--- Testing Excel ---")
    #         wb = openpyxl.Workbook()
    #         sheet = wb.active
    #         sheet['A1'] = "Hello"
    #         sheet['B1'] = "World"
    #         sheet.cell(row=2, column=1, value="This is a test.")
    #         wb.save("test.xlsx")
    #         excel_text = read_excel("test.xlsx")
    #         print(f"Extracted Excel Text:\n{excel_text}")
    #         os.remove("test.xlsx")
    #     except DocumentProcessingError as e:
    #         print(f"Excel Error: {e}")
    #     except Exception as e:
    #         print(f"Excel test setup error: {e}")
    # else:
    print("Excel test skipped (requires openpyxl and a valid test.xlsx file).")


    # Test Word (requires a dummy file named test.docx)
    # if DocxDocument:
    #     try:
    #         print("\n--- Testing Word ---")
    #         doc = DocxDocument()
    #         doc.add_paragraph("First paragraph.")
    #         doc.add_paragraph("Second paragraph with some more text.")
    #         doc.save("test.docx")
    #         word_text = read_word("test.docx")
    #         print(f"Extracted Word Text:\n{word_text}")
    #         os.remove("test.docx")
    #     except DocumentProcessingError as e:
    #         print(f"Word Error: {e}")
    #     except Exception as e:
    #         print(f"Word test setup error: {e}")
    # else:
    print("Word test skipped (requires python-docx and a valid test.docx file).")
